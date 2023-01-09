# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.

from odoo import api, fields, models, _
from odoo.exceptions import UserError, ValidationError
from odoo.http import request
from datetime import datetime, date
import base64
from io import BytesIO
# from docx import Document
import requests
# import requests_pkcs12
from OpenSSL import crypto
# from docx2pdf import convert
from docx.shared import Inches, Pt
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
import pdfminer
import json
import os

import smtplib
import imaplib
# import email
# from email.header import decode_header
# from email.header import make_header

from email import encoders
# from email.message import Message
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import time
from docxtpl import DocxTemplate
import pyotp
from pytz import timezone

from PyPDF2 import PdfFileWriter, PdfFileReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from cryptography import x509
from cryptography.hazmat.backends import default_backend
import qrcode

import logging

from pythainlp.util import thai_strftime

_logger = logging.getLogger(__name__)

SIGN_CH = (
    ("sign", "ลงลายมือชื่อ"),
    ("img", "เลือกภาพลายเซ็นต์")
)
TOTP = pyotp.TOTP('base32secret3232', interval=300)
HOTP = pyotp.HOTP('base32secret3232')


class MakeApprovalSignatureWizardInherit(models.TransientModel):
    _inherit = "make.approval.wizard"

    # memo = fields.Text('Memo', required=False)
    sign = fields.Binary(
        string='ลงลายมือชื่อ',
        required=False)
    img = fields.Binary(
        string='ภาพลายเซ็นต์',
        required=False,
        default=lambda self: self.env['hr.employee'].search([('user_id', "=", self.env.uid)], limit=1).sign_img
    )
    sign_method = fields.Selection(
        SIGN_CH,
        string='เลือกวิธีลงลายเซ็นต์',
        default='sign'
    )
    sign_otp = fields.Char(
        string='รหัส OTP (2FA)',
    )
    sign_otp_temp = fields.Char(
        default=False
    )
    sign_otp_counter = fields.Integer(
        default=0
    )

    # cert_passphrase = fields.Char(
    #     string='รหัสผ่าน'
    # )

    def _get_dsign_auth(self):
        sign_auth = False
        if 'dsign_auth' in request.session:
            sign_auth = True
        return sign_auth

    def _get_internal_document_type(self):
        internal = self.env['document.internal.main'].search([
            ('id', '=', self.env.context.get('active_id'))
        ], limit=1)
        if internal:
            return internal.document_type

    sign_auth = fields.Boolean(
        default=_get_dsign_auth,
    )
    action_request_otp_button = fields.Boolean(
        string='ขอรหัส OTP',
        default=False,
    )
    doc_type = fields.Char(
        default=_get_internal_document_type
    )

    # data to document
    def from_data(self, current_id):
        cur_path = os.path.dirname(os.path.abspath(__file__))
        docx_path = cur_path + '/dSign/depa_template.docx'
        rec_ids = self.env['document.internal.main'].search([('id', 'in', [current_id])])
        document = DocxTemplate(docx_path)
        # document = Document()
        # document.add_heading('หนังสือภายใน', 9)
        document.add_paragraph('')
        style = document.styles['Normal']
        font = style.font
        font.name = 'TH SarabunPSK'
        font.size = Pt(16)
        section = document.sections[0]
        footer = section.footer
        for doc_ids in rec_ids:
            circular_letter = str(doc_ids.circular_letter)
            name = str(doc_ids.name)
            date = str(doc_ids.date_document)
            real_name = str(doc_ids.name_real)
            real_date = str(doc_ids.date_document_real)
            depart = str(doc_ids.department_name)
            speed = str(doc_ids.speed)
            secret = str(doc_ids.secret)
            subject = str(doc_ids.subject)
            dear = str(doc_ids.dear)
            employee = str(doc_ids.empolyee_name.name)
            sign = str(doc_ids.sign.name)
            for_document = str(doc_ids.for_document)
            other = str(doc_ids.for_document_other)
            note = str(doc_ids.note)
            material = str(doc_ids.material)

            fist_para = document.add_paragraph('หนังสือเวียน ')
            fist_para.add_run(circular_letter)
            sec_para = document.add_paragraph('Name ')
            sec_para.add_run(name)
            thir_para = document.add_paragraph('Date ')
            thir_para.add_run(date)
            four_para = document.add_paragraph('Real Name ')
            four_para.add_run(real_name)
            fift_para = document.add_paragraph('Real Date ')
            fift_para.add_run(real_date)
            six_para = document.add_paragraph('Department ')
            six_para.add_run(depart)
            sev_para = document.add_paragraph('Speed ')
            sev_para.add_run(speed)
            eigh_para = document.add_paragraph('Secret ')
            eigh_para.add_run(secret)
            nine_para = document.add_paragraph('Subject ')
            nine_para.add_run(subject)
            ten_para = document.add_paragraph('Dear ')
            ten_para.add_run(dear)
            sixteen_para = document.add_paragraph('Employee ')
            sixteen_para.add_run(employee)
            ele_para = document.add_paragraph('Sign ')
            ele_para.add_run(sign)
            twel_para = document.add_paragraph('For Document ')
            twel_para.add_run(for_document)
            thirteen_para = document.add_paragraph('Material ')
            thirteen_para.add_run(material)
            fourteen_para = document.add_paragraph('Other ')
            fourteen_para.add_run(other)
            fifteen_para = document.add_paragraph('Note ')
            fifteen_para.add_run(note)
            document.add_paragraph('')
        fontx = footer.paragraphs[0].add_run(
            'เลขที่อ้างอิง 1234 IN-xxxx สําานักงาน กสทช. และ สศด. TH-e-GIF e-CMS version 2.0 on Cloud').font
        fontx.name = 'TH SarabunPSK'
        fontx.size = Pt(10)
        # document.save('test.docx')

        fp = BytesIO()
        document.save(fp)
        # return fp.getvalue()
        fp.seek(0)
        data = fp.read()
        fp.close()
        return data

    def arabic_to_thai_number(self, number):
        thai = {
            '0': '๐',
            '1': '๑',
            '2': '๒',
            '3': '๓',
            '4': '๔',
            '5': '๕',
            '6': '๖',
            '7': '๗',
            '8': '๘',
            '9': '๙'
        }
        thai_number = ''
        for i in number:
            if i.isnumeric():
                thai_number += thai[i]
            else:
                thai_number += i
        return thai_number

    def get_thai_month(self, number):
        month = {
            '1': 'มกราคม',
            '2': 'กุมภาพันธ์',
            '3': 'มีนาคม',
            '4': 'เมษายน',
            '5': 'พฤษภาคม',
            '6': 'มิถุนายน',
            '7': 'กรกฎาคม',
            '8': 'สิงหาคม',
            '9': 'กันยายน',
            '10': 'ตุลาคม',
            '11': 'พฤศจิกายน',
            '12': 'ธันวาคม',
        }
        return month[number]

    def insertDocNumber(self, doc_id, pdf_base64, invitation_name=None):
        doc = self.env['document.internal.main'].browse(int(doc_id))
        base64_input = pdf_base64
        base64_output = ''
        input_file = os.path.dirname(os.path.abspath(__file__)) + '/temp/temp_pdf_input.pdf'
        output_file = os.path.dirname(os.path.abspath(__file__)) + '/temp/temp_pdf_output.pdf'
        font_path = os.path.dirname(os.path.abspath(__file__)) + '/font/THSarabunNew.ttf'
        font_bold_path = os.path.dirname(os.path.abspath(__file__)) + '/font/THSarabunNewBold.ttf'
        qrcode_dir = os.path.dirname(os.path.abspath(__file__)) + '/temp/qr_code.jpg'

        # convert base64 to input pdf file
        with open(input_file, 'wb') as pdfFile:
            pdfFile.write(base64.b64decode(str(base64_input)))

        # set custom font
        pdfmetrics.registerFont(TTFont('THSarabunNew', font_path))
        pdfmetrics.registerFont(TTFont('THSarabunNewBold', font_bold_path))
        # create steam io packet
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        # can.setFillColorRGB(1, 0, 0)
        can.setFont('THSarabunNew', 16)

        if doc.document_type_select in ["หนังสือภายนอก+หนังสือรับรอง", "หนังสือรับรอง"]:
            if doc.name_real:
                # can.drawString(100, 700, doc.name_real)
                name_real = "ที่ " + str(self.arabic_to_thai_number(doc.name_real))
                if doc.specific_document_no:
                    name_real = "ที่ " + str(self.arabic_to_thai_number(doc.specific_document_no))
                can.drawString(85, 709, name_real)
            else:
                can.drawString(85, 700, " ")
            # can.save()

            if doc.document_type_select not in ["หนังสือรับรอง"]:
                if doc.date_document_real:
                    dates = doc.date_document_real
                    date_thai = self.arabic_to_thai_number(str(dates.day)) + " " + self.get_thai_month(
                        str(dates.month)) + " " + self.arabic_to_thai_number(str(dates.year + 543))
                    if doc.specific_date:
                        specific_date = doc.specific_date
                        date_thai = self.arabic_to_thai_number(str(specific_date.day)) + " " + self.get_thai_month(
                            str(specific_date.month)) + " " + self.arabic_to_thai_number(str(specific_date.year + 543))
                    can.drawString(320, 685, date_thai)
        elif "คำสั่ง" in doc.document_type_select:
            if doc.name_real:
                name_real = "ที่ " + str(self.arabic_to_thai_number(doc.name_real))
                if doc.specific_document_no:
                    name_real = "ที่ " + str(self.arabic_to_thai_number(doc.specific_document_no))
                co_x = 289 - len(name_real)
                can.setFont('THSarabunNewBold', 16)
                can.drawString(co_x, 675, name_real)
                can.setFont('THSarabunNew', 16)
        else:
            can.drawString(85, 709, " ")

        date_pdf = False
        date_page = 0

        # กรณีเป็น คำสั่ง/ระเบียบ/ข้อบังคับ/ประกาศ จะต้องใส่วันที่ต่อจากคำว่า 'ประกาศ ณ วันที่' หรือ 'สั่ง ณ วันที่'
        if doc.document_type_select in ['ประกาศ', 'ระเบียบ', 'ข้อบังคับ', 'หนังสือรับรอง'] or 'คำสั่ง' in doc.document_type_select:
            if doc.date_document_real:
                dates = doc.date_document_real
                date_thai = self.arabic_to_thai_number(str(dates.day)) + " " + self.get_thai_month(
                    str(dates.month)) + " พ.ศ. " + self.arabic_to_thai_number(str(dates.year + 543))
                if doc.specific_date:
                    specific_date = doc.specific_date
                    date_thai = self.arabic_to_thai_number(str(specific_date.day)) + " " + self.get_thai_month(
                        str(specific_date.month)) + " พ.ศ. " + self.arabic_to_thai_number(str(specific_date.year + 543))
                date_packet = BytesIO()
                date_can = canvas.Canvas(date_packet, pagesize=A4)
                date_can.setFont('THSarabunNew', 16)

                fp = open(input_file, 'rb')
                parser = PDFParser(fp)
                document = PDFDocument(parser)
                rsrcmgr = PDFResourceManager()
                device = PDFDevice(rsrcmgr)
                laparams = LAParams()
                device = PDFPageAggregator(rsrcmgr, laparams=laparams)
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                document_type = doc.document_type_select

                coordinate_x = coordinate_y = 0

                for pageNumber, page in enumerate(PDFPage.create_pages(document), 1):
                    interpreter.process_page(page)
                    layout = device.get_result()
                    left, bottom, text = self.parse_pdf_obj_date(layout._objs, document_type)
                    if left > 0 and bottom > 0:
                        date_page = pageNumber
                        coordinate_x = left
                        coordinate_y = bottom
                        # break
                # raise ValidationError(f"{date_page} | {coordinate_x} | {coordinate_y}")

                if coordinate_x > 0 and coordinate_y > 0:
                    if 'คำสั่ง' in doc.document_type_select:
                        adj_x = 43
                    elif doc.document_type_select == 'หนังสือรับรอง':
                        adj_x = 58
                    else:
                        adj_x = 73
                    adj_y = 4
                    date_can.drawString(coordinate_x + adj_x, coordinate_y + adj_y, str(date_thai))
                    date_can.save()
                    date_packet.seek(0)
                    date_pdf = PdfFileReader(date_packet)


        can.save()

        # move to the beginning of the StringIO buffer
        packet.seek(0)

        # create a new PDF with Reportlab
        new_pdf = PdfFileReader(packet)

        # read your existing PDF
        existing_pdf = PdfFileReader(open(input_file, "rb"))
        output = PdfFileWriter()

        # add the "watermark" (which is the new pdf) on the existing page
        # page = existing_pdf.getPage(0)
        # page.mergePage(new_pdf.getPage(0))
        # output.addPage(page)

        # if qrcode_dir:
        #     qr_packet = BytesIO()
        #     qr_can = canvas.Canvas(qr_packet, pagesize=A4)
        #     qr_can.drawImage(qrcode_dir, 85, 20, width=50, height=50, mask='auto')
        #     qr_can.setFont('THSarabunNew', 8)
        #     qr_can.drawString(90, 10, 'Digital Signature')
        #     qr_can.save()
        #     qr_packet.seek(0)
        #     qr_pdf = PdfFileReader(qr_packet)

        invitation_pdf = False

        if invitation_name is not None:
            invitation_packet = BytesIO()
            invitation_can = canvas.Canvas(invitation_packet, pagesize=A4)
            invitation_can.setFont('THSarabunNew', 16)

            fp = open(input_file, 'rb')
            parser = PDFParser(fp)
            document = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            device = PDFDevice(rsrcmgr)
            laparams = LAParams()
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            document_type = doc.document_type_select

            coordinate_x = coordinate_y = 0
            coordinate_text = ""

            for pageNumber, page in enumerate(PDFPage.create_pages(document), 1):
                interpreter.process_page(page)
                layout = device.get_result()
                left, bottom, text = self.parse_pdf_obj_invitation(layout._objs, document_type)
                if left > 0 and bottom > 0:
                    coordinate_text = text
                    coordinate_x = left
                    coordinate_y = bottom
                    break

            if coordinate_x > 0 and coordinate_y > 0:
                adj_x = 0
                if "เรียน" not in coordinate_text:
                    adj_x = 30
                invitation_can.drawString(coordinate_x + 30 + adj_x, coordinate_y + 4, str(invitation_name))
                # raise ValidationError(invitation_name)
                invitation_can.save()
                invitation_packet.seek(0)
                invitation_pdf = PdfFileReader(invitation_packet)

        for i in range(existing_pdf.getNumPages()):
            # if i == 0 and i == (existing_pdf.getNumPages() - 1):
            #     page = existing_pdf.getPage(0)
            #     page.mergePage(new_pdf.getPage(0))
            #     # page.mergePage(qr_pdf.getPage(0))
            #     if invitation_pdf:
            #         page.mergePage(invitation_pdf.getPage(0))
            #     if date_pdf:
            #         page.mergePage(date_pdf.getPage(0))
            #     output.addPage(page)

            # elif i == 0:
            #     page = existing_pdf.getPage(0)
            #     page.mergePage(new_pdf.getPage(0))
            #     if invitation_pdf:
            #         page.mergePage(invitation_pdf.getPage(0))
            #     output.addPage(page)

            # elif i == (date_page - 1):
            #     page = existing_pdf.getPage(date_page - 1)
            #     if date_pdf:
            #         page.mergePage(date_pdf.getPage(0))
            #     # page.mergePage(qr_pdf.getPage(0))
            #     output.addPage(page)

            # else:
            #     output.addPage(existing_pdf.getPage(i))

            page = existing_pdf.getPage(i)
            if i == 0:
                page.mergePage(new_pdf.getPage(0)) # เลขที่หนังสือ
                if invitation_pdf:
                    page.mergePage(invitation_pdf.getPage(0)) # รายชื่อแจ้งท้าย
            if i == (date_page - 1):
                if date_pdf:
                    page.mergePage(date_pdf.getPage(0)) # วันที่หนังสือ
            output.addPage(page)

        # finally, write "output" to a real file
        outputStream = open(output_file, "wb")
        output.write(outputStream)
        outputStream.close()

        # convert input pdf file to base64 encoded
        with open(output_file, "rb") as pdf_file:
            base64_output = base64.b64encode(pdf_file.read())

        # remove temp pdf file
        try:
            os.remove(input_file)
            os.remove(output_file)
        except OSError as e:
            raise ValidationError(_(e))

        return base64_output

    def generate_certificate(self, cert_path):
        # cur_path = os.path.dirname(os.path.abspath(__file__))
        # pkcs12_path = cur_path + '/dSign/uat-D-SignSp.p12'
        # pkcs12_path = cur_path + '/dSign/DEPA_Service_Provider_Production.p12'
        # pkcs12_pass = cur_path + '/dSign/pwd-for-pkcs12.txt'
        hasSettings = self.env['depa_signature_setting'].search([], limit=1)
        if not hasSettings:
            raise ValidationError(
                _('กรุณาตั้งค่า Digital Signature ไฟล์ให้เสร็จสมบูรณ์'))

        if not hasSettings.cert_path_file:
            raise ValidationError(
                _('กรุณาตั้งค่าไฟล์สำหรับเชื่อมต่อ Digital Signature ให้เสร็จสมบูรณ์'))

        pkcs12_path = hasSettings.cert_path_file

        passw = base64.b64decode(hasSettings.cert_passphrase.encode('utf-8')).decode('utf-8')
        with open(pkcs12_path, 'rb') as f:
            fp = f.read()

        # with open(pkcs12_pass, 'rb') as f:
        #     passw = f.read()

        # if self.sign_auth:
        #     passw = request.session['cert_passphrase']
        # else:
        #     passw = self.cert_passphrase
        #     self.cert_passphrase = False

        try:
            p12 = crypto.load_pkcs12(fp, passw)
            # if not self.sign_auth:
            #     request.session['cert_passphrase'] = passw
        except:
            raise ValidationError(
                _('รหัสผ่านในการเข้าถึง CA ของคุณไม่ถูกต้อง'))

        temp_dir = os.path.dirname(os.path.abspath(__file__)) + '/temp'
        try:
            # os.chmod(temp_dir, 0o777)
            # PEM formatted certificate
            c = crypto.dump_certificate(crypto.FILETYPE_PEM, p12.get_certificate())
            fcert = open(cert_path[0], 'wb')
            fcert.write(c)
            fcert.close()

            # PEM formatted private key
            k = crypto.dump_privatekey(crypto.FILETYPE_PEM, p12.get_privatekey())
            fkey = open(cert_path[1], 'wb')
            fkey.write(k)
            fkey.close()

            # encrypt public certificate
            # raise ValidationError(c)
            # string_cert = str(c).replace('\\n', '\n')
            qr_cert = x509.load_pem_x509_certificate(c, default_backend())

            # create qr code from cert string
            if qr_cert:
                qrcode_dir = os.path.dirname(os.path.abspath(__file__)) + '/temp/qr_code.jpg'
                qr = qrcode.QRCode(
                    version=1,
                    error_correction=qrcode.constants.ERROR_CORRECT_H,
                    box_size=10,
                    border=4,
                )
                qr.add_data(qr_cert.serial_number)
                qr.make(fit=True)
                img = qr.make_image(fill_color="white", back_color="black").convert('RGB')
                img.save(qrcode_dir)

        except Exception as e:
            raise ValidationError(e)
            # raise ValidationError(
            #     _('พบปัญหาระหว่างการประมวลผล Temp File - 1'))
        finally:
            os.chmod(temp_dir, 0o755)

    def remove_certificate(self, cert_path):
        status = True

        for file in cert_path:
            if os.path.exists(file):
                os.remove(file)
                continue
            else:
                status = False
                break
        return status

    def convert_image_signature(self, file_path):
        encoded_string = ""
        if file_path != "":
            with open(file_path, "rb") as image_file:
                encoded_string = base64.b64encode(image_file.read())
        return encoded_string

    def convert_document(self, doc_id, docx_file, invitation_name=None, converted_document=None):
        # covert docx to pdf
        # url = "http://202.139.197.40:7777/wordtopdf"
        if converted_document is not None:
            response_code = converted_document['code']
            response_data = converted_document['data']
        else:
            url = "https://api.depa.or.th/wordtopdf2"
            file_name = "docx_test_file.docx"
            data_base64 = ""
            data_status = False
            action = {"docxFile": (file_name, docx_file)}
            headers = { 'User-Agent': 'Mozilla/5.0 (Windows NT 6.0; WOW64; rv:24.0) Gecko/20100101 Firefox/24.0' }
            resp = requests.post(url, files=action, headers=headers)
            response_code = resp.status_code
            response_data = resp.text

        if response_code == 200:
            data_status = True
            json_data = json.loads(response_data)
            data_base64 = self.insertDocNumber(doc_id, json_data['base64'], invitation_name=invitation_name)

        return data_status, data_base64

    def convert_plain_document(self, docx_file):
        url = "https://api.depa.or.th/wordtopdf2"
        file_name = "docx_test_file.docx"
        data_base64 = ""
        data_status = False
        action = {"docxFile": (file_name, docx_file)}
        headers = { 'User-Agent': 'Mozilla/5.0 (Windows NT 6.0; WOW64; rv:24.0) Gecko/20100101 Firefox/24.0' }
        resp = requests.post(url, files=action, headers=headers)
        response = {}
        response['code'] = resp.status_code
        response['data'] = resp.text

        return response

    def get_account_approver_sign(self, position):
        approver_data = self.env['account_approver_sign'].search([
            ('position', '=', position),
            ('is_used', '=', True)
        ], limit=1)
        if approver_data:
            sign_img = approver_data['employee_id'].sign_img
            cad_pass = self.env['depa_signature_setting_lines'].search([
                ('employee_id', '=', approver_data['employee_id'].id),
                ('disable', '=', False),
                ('active', '=', True)
            ], limit=1).cad_password
            if not cad_pass:
                raise ValidationError(f"ไม่พบ CAD PASSWORD ของ {approver_data.name}")
            # sign_img = self.env['hr.employee'].search([('id', "=", emp_id)], limit=1).sign_img
            return sign_img, cad_pass
        return False


    def parse_pdf_obj(self, lt_objs, document_type):
        left_pos = 0
        bottom_pos = 0
        text_line = ""

        # หนังสือภายนอก+หนังสือรับรอง
        # ประกาศ
        # ระเบียบ
        # ข้อบังคับ
        for obj in lt_objs:
            # if it's a textbox, print text and location
            # raise ValidationError(lt_objs)
            if isinstance(obj, pdfminer.layout.LTTextBoxHorizontal):
                if document_type in ['หนังสือภายนอก+หนังสือรับรอง','หนังสือรับรอง']:
                    if ('ขอ' in obj.get_text() and 'แสดง' in obj.get_text() and 'ความ' in obj.get_text() and 'นับ' in obj.get_text() and 'ถือ' in obj.get_text()) or ('Yours sincerely' in obj.get_text()) or ('ณ วันที' in obj.get_text()):
                        # print(obj.get_text())
                        left_pos = obj.bbox[0]
                        bottom_pos = obj.bbox[1]
                        text_line = obj.get_text().replace('\n', '_')
                elif document_type in ['ประกาศ', 'ระเบียบ', 'ข้อบังคับ']:
                    if ('ประกาศ' in obj.get_text() and 'ณ' in obj.get_text() and 'วันที' in obj.get_text()):
                    # if 'ประกาศ ณ วันที' in obj.get_text():
                        left_pos = obj.bbox[0]
                        bottom_pos = obj.bbox[1]
                        text_line = obj.get_text().replace('\n', '_')
                elif "คำสั่ง" in document_type:
                    if ('ณ' in obj.get_text() and 'วันที' in obj.get_text()):
                    # if 'ณ วันที' in obj.get_text():
                        left_pos = obj.bbox[0]
                        bottom_pos = obj.bbox[1]
                        text_line = obj.get_text().replace('\n', '_')
                elif document_type == 'หนังสือรับรอง':
                    if ('ณ' in obj.get_text() and 'วันที' in obj.get_text()):
                    # if 'ณ วันที' in obj.get_text():
                        left_pos = obj.bbox[0]
                        bottom_pos = obj.bbox[1]
                        text_line = obj.get_text().replace('\n', '_')
            # if it's a container, recurse
            # elif isinstance(obj, pdfminer.layout.LTFigure):
            #     self.parse_pdf_obj(obj._objs, document_type)
        return left_pos, bottom_pos, text_line

    def parse_pdf_obj_invitation(self, lt_objs, document_type):
        left_pos = 0
        bottom_pos = 0
        text_line = ""

        # หนังสือภายนอก+หนังสือรับรอง
        # ประกาศ
        # ระเบียบ
        # ข้อบังคับ
        # texts = []
        for obj in lt_objs:
            # if it's a textbox, print text and location
            if isinstance(obj, pdfminer.layout.LTTextBoxHorizontal):
                # texts.append(str(obj.get_text()))
                # if 'เรียน' in obj.get_text() and '\n' in obj.get_text():
                if 'เรียน ' in obj.get_text() or 'Excellency,' in obj.get_text():
                    left_pos = obj.bbox[0]
                    bottom_pos = obj.bbox[1]
                    text_line = obj.get_text().replace('\n', '_')
                    break
            # if it's a container, recurse
            elif isinstance(obj, pdfminer.layout.LTFigure):
                self.parse_pdf_obj(obj._objs, document_type)
        # raise ValidationError(texts)
        return left_pos, bottom_pos, text_line

    def parse_pdf_obj_date(self, lt_objs, document_type):
        left_pos = 0
        bottom_pos = 0
        text_line = ""

        # หนังสือภายนอก+หนังสือรับรอง
        # ประกาศ
        # ระเบียบ
        # ข้อบังคับ
        # texts = []
        # raise ValidationError(lt_objs)
        for obj in lt_objs:
            # if it's a textbox, print text and location
            if isinstance(obj, pdfminer.layout.LTTextBoxHorizontal):
                # texts.append(obj.get_text())
                if document_type in ['ประกาศ', 'ระเบียบ', 'ข้อบังคับ']:
                    if "ประกาศ" in obj.get_text() and "ณ" in obj.get_text() and "วันที" in obj.get_text():
                    # if 'ประกาศ ณ วันที' in obj.get_text():
                        left_pos = obj.bbox[0]
                        bottom_pos = obj.bbox[1]
                        text_line = obj.get_text().replace('\n', '_')
                        # break
                elif "คำสั่ง" in document_type:
                    if "ณ" in obj.get_text() and "วันที" in obj.get_text():
                    # if 'ง ณ วันที' in obj.get_text():
                        left_pos = obj.bbox[0]
                        bottom_pos = obj.bbox[1]
                        text_line = obj.get_text().replace('\n', '_')
                elif "หนังสือรับรอง" in document_type:
                    if "ณ" in obj.get_text() and "วันที" in obj.get_text():
                    # if 'ณ วันที' in obj.get_text():
                        left_pos = obj.bbox[0]
                        bottom_pos = obj.bbox[1]
                        text_line = obj.get_text().replace('\n', '_')

            # if it's a container, recurse
            # elif isinstance(obj, pdfminer.layout.LTFigure):
            #     raise ValidationError(obj)
                # self.parse_pdf_obj(obj._objs, document_type)
        # raise ValidationError(texts)
        return left_pos, bottom_pos, text_line

    def get_coordinate_signature(self, data_encoded, document_type):
        pdf_file = os.path.dirname(os.path.abspath(__file__)) + '/temp/temp_pdf.pdf'
        temp_dir = os.path.dirname(os.path.abspath(__file__)) + '/temp'

        try:
            # os.chmod(temp_dir, 0o777)
            with open(pdf_file, 'wb') as pdfFile:
                pdfFile.write(base64.b64decode(data_encoded))
        except:
            raise ValidationError(_('พบปัญหาระหว่างการประมวลผล Temp File - 2'))
        finally:
            os.chmod(temp_dir, 0o755)

        fp = open(pdf_file, 'rb')
        parser = PDFParser(fp)
        document = PDFDocument(parser)
        if not document.is_extractable:
            raise PDFTextExtractionNotAllowed
        rsrcmgr = PDFResourceManager()
        device = PDFDevice(rsrcmgr)
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        coordinate_x = coordinate_y = coordinate_page = 0

        for pageNumber, page in enumerate(PDFPage.create_pages(document), 1):
            interpreter.process_page(page)
            layout = device.get_result()
            left, bottom, text = self.parse_pdf_obj(layout._objs, document_type)
            if left > 0 and bottom > 0:
                coordinate_x = left
                coordinate_y = bottom
                coordinate_page = pageNumber
                # break

        if os.path.exists(pdf_file):
            os.remove(pdf_file)

        return coordinate_x, coordinate_y, coordinate_page

    def sign_document(self, doc_id, signature_file=None, cad_password=None, invitation_name=None, converted_document=None, specific_x=None, specific_y=None, pdf_file=None):

        internal_doc = self.env['document.internal.main'].browse(int(doc_id))

        cert_file = (
            os.path.dirname(os.path.abspath(__file__)) + '/temp/myCert.crt',
            os.path.dirname(os.path.abspath(__file__)) + '/temp/myKey.key'
        )
        self.generate_certificate(cert_file)

        if pdf_file is None:
            convert_status, data_encoded = self.convert_document(
                # self.from_data(doc_id)
                # self.env['saraban_form'].from_data(doc_id) #Convert Custom Format
                internal_doc.id,
                base64.b64decode(internal_doc.file_for_signature),  # Convert From File
                invitation_name=invitation_name,
                converted_document=converted_document
            )

            if not convert_status:
                if invitation_name is not None:
                    raise ValidationError(_('เกิดข้อผิดพลาดในการสร้างและประมวลผลเอกสาร กรุณาลองใหม่อีกครั้ง'+str(invitation_name)))
                else:
                    raise ValidationError(_('เกิดข้อผิดพลาดในการสร้างและประมวลผลเอกสาร กรุณาลองใหม่อีกครั้ง'))
        else:
            convert_status = True
            data_encoded = pdf_file
        

        pdf_data = ""
        sign_status = False
        if convert_status:
            if signature_file is None:
                raise ValidationError(
                    _('กรุณาระบุภาพลายเซ็นต์ของคุณให้ถูกต้อง'))

            if specific_x is None or specific_y is None:
                x, y, page_number = self.get_coordinate_signature(data_encoded=data_encoded,
                                                                document_type=internal_doc.document_type)
            else:
                x = specific_x / 1000
                y = specific_y / 1000
                page_number = 1

            if x > 0 and y > 0:
                # cad_path = os.path.dirname(os.path.abspath(__file__)) + '/dSign/cad-data.txt'
                # cad_path = os.path.dirname(os.path.abspath(__file__)) + '/dSign/Credential_Data_Depa.cad'

                # cert_file = (
                #     os.path.dirname(os.path.abspath(__file__)) + '/temp/myCert.crt',
                #     os.path.dirname(os.path.abspath(__file__)) + '/temp/myKey.key'
                # )
                # self.generate_certificate(cert_file)

                # with open(cad_path, "r") as f:
                #     cad = f.read().strip()

                sign_image = signature_file
            
                if specific_x is None or specific_y is None:
                    if internal_doc['document_type_select'] == 'หนังสือภายนอก+หนังสือรับรอง':
                        adj_x = 180
                    elif internal_doc['document_type_select'] in ['ประกาศ', 'ระเบียบ', 'ข้อบังคับ']:
                        adj_x = 340
                    elif 'คำสั่ง' in internal_doc['document_type_select'] or internal_doc['document_type_select'] == 'หนังสือรับรอง':
                        adj_x = 300

                    if (internal_doc['document_type_select'] in ['ประกาศ', 'ระเบียบ', 'ข้อบังคับ']) or 'คำสั่ง' in internal_doc['document_type_select']:
                        adj_y = 0
                    elif internal_doc['document_type_select'] == 'หนังสือรับรอง':
                        adj_y = -10
                    else:
                        adj_y = 10

                    delta_y = (y - 500) / 5
                    x = int("%6d" % ((x + adj_x))) / 1000
                    y = int("%6d" % ((y + adj_y + delta_y))) / 1000

                try:
                    data_encoded = data_encoded.decode('utf-8')
                except:
                    pass

                sign_payloads = {
                    # "pdfData": data_encoded.decode('utf-8'),
                    "pdfData": data_encoded,
                    "cadData": str(cad_password),
                    "reason": "ลงนามเอกสาร",
                    "location": "TH",
                    "certifyLevel": "NON-CERTIFY",
                    "overwriteOriginal": True,
                    "visibleSignature": "Graphics",
                    "visibleSignaturePage": page_number,
                    "visibleSignatureRectangle": str(x) + ", " + str(y) + ", 0.2, 0.1",
                    "visibleSignatureImagePath": sign_image.decode()
                }
                # print("Payloads: ", sign_payloads)
                hasSettings = self.env['depa_signature_setting'].search([], limit=1)
                if not hasSettings:
                    raise ValidationError(
                        _('กรุณาตั้งค่า Digital Signature ไฟล์ให้เสร็จสมบูรณ์'))

                if not hasSettings.url_endpoint:
                    raise ValidationError(
                        _('กรุณากำหนด API Endpoint URL ให้ถูกต้อง'))

                # sign_url = 'https://uat-sign.one.th/webservice/api/v2/signing/pdfSigning-V3'
                # sign_url = 'https://sign.one.th/webservice/api/v2/signing/pdfSigning-V3'

                sign_url = hasSettings.url_endpoint
                try:
                    sign_resp = requests.post(sign_url, json=sign_payloads, cert=cert_file)
                except:
                    raise ValidationError(
                        _('พบปัญหาระหว่างการ Signing Document'))
                finally:
                    self.remove_certificate(cert_file)

                if sign_resp.status_code == 200:
                    resp_data = sign_resp.json()
                    # print(resp_data)
                    if resp_data['pdfData'] is not None:
                        pdf_data = resp_data['pdfData']
                        sign_status = True

        return sign_status, pdf_data

    def getDocNumber(self, doc_number_name):
        last_number = None
        if doc_number_name.find('/') != -1:
            name = doc_number_name.split('/')
            last_number = name[len(name) - 1].strip()
        return last_number

    def getExtenstionName(self, filename):
        ext_name = None
        if filename.find('.') != -1:
            name = filename.split('.')
            ext_name = name[len(name) - 1].strip()
        return ext_name

    def set_draft_email(self, doc_id, pdf_signed=None):
        doc = self.env['document.internal.main'].search([('id', '=', [doc_id])])
        # credential = os.path.dirname(os.path.abspath(__file__)) + '/dSign/credential_sarabun.txt'

        doc_speed = dict((
            ("1", "ปกติ"),
            ("2", "ด่วน"),
            ("3", "ด่วนมาก"),
            ("4", "ด่วนที่สุด")
        ))

        prefix_sub = ''
        if doc.speed != "1":
            prefix_sub = f'{doc_speed[doc.speed]} : '

        # content body
        subject = f"{prefix_sub}{doc.internal_title}"
        msgBody = "<br>".join(doc.material.split("\n"))
        for_doc = f"จึงเรียนมา{doc.for_document}"
        if doc.for_document == 'other':
            for_doc = doc.for_document_other
        msgText = f"""
        <html>
          <head></head>
          <body>
            <p>
                เรียน {doc.dear_select}<br><br>
                {msgBody}<br><br>
                {for_doc}<br><br>

                <pre>สำนักงานส่งเสริมเศรษฐกิจดิจิทัล<br>{doc.suffix}<br>โทร. 020 262 333<br>โทรสาร 020 262 333 กด 9 </pre><br><br>
                ___________________________________________________________________________<br>
                หากท่านได้รับอีเมลนี้แล้วกรุณาแจ้งการได้รับกลับมายังที่อยู่อีเมลนี้ด้วย จะขอบคุณยิ่ง<br><br>

                อีเมล (และ/หรือเอกสารแนบ) นี้เป็นข้อมูลที่ความลับและอาจเป็นข้อมูลที่เป็นเอกสิทธิ์เฉพาะบุคคล การนำข้อมูลดังกล่าวไปใช้หรือเปิดเผยให้บุคคลอื่นใดล่วงรู้ เป็นการกระทำที่ไม่ได้รับอนุญาต หากท่านมิได้เป็นบุคคลที่อีเมลฉบับนี้ระบุถึงแล้ว กรุณาลบอีเมลนี้ออกจากคอมพิวเตอร์ที่ท่านได้รับ ทั้งนี้ ขอขอบคุณหากท่านได้แจ้งผู้ส่งถึงการจัดส่งอีเมลผิดพลาดด้วย<br><br>

                This e-mail (and/or attachments) is confidential and may be privileged. Use or disclosure of it by anyone other than a designated addressee is unauthorized. If you are not an intended recipient, please delete this e-mail from the computer on which you received it. We thank you for notifying us immediately.
            </p>
          </body>
        </html>
        """
        # print(doc.attachment_ids)

        # Connect to inbox
        SMTP_SERVER = "imap.gmail.com"
        # SMTP_PORT = 993

        # email_address = "saraban@depa.or.th"
        # with open(credential, "r") as f:
        #     email_password = f.read().strip()

        hasSettings = self.env['depa_signature_setting'].search([], limit=1)
        if not hasSettings:
            raise ValidationError(
                _('กรุณาตั้งค่า email สำหรับฉบับร่าง'))
        email_address = hasSettings.email_to_draft
        email_password = base64.b64decode(hasSettings.email_credential.encode('utf-8')).decode('utf-8')

        imap_server = imaplib.IMAP4_SSL(SMTP_SERVER)
        imap_server.login(email_address, email_password)
        # Select mailbox
        imap_server.select("[Gmail]/Drafts")
        # Create message
        new_message = MIMEMultipart()
        new_message["From"] = "Saraban Depa <saraban@depa.or.th>"
        new_message["To"] = "Name of Recipient <recpient@mydomain.com>"
        new_message["Subject"] = subject
        html = """
        <html>
          <head></head>
          <body>
            <p>Hi!<br>
               How are you?<br>
               Here is the <a href="http://www.python.org">link</a> you wanted.
            </p>
          </body>
        </html>
        """

        msgText = MIMEText(msgText, 'html')
        new_message.attach(msgText)

        filename_format = 'สำเนาหนังสือลงนาม (Digital Signature)'
        if doc.name_real:
            reformat = '_depa'
            if 'สศด.' in doc.name_real:
                reformat = f'_depa{doc.name_real[4:].replace(" /", "_")}'
            filename_format = str(date.today().year + 543) + reformat
            # filename_format = str(date.today().year + 543) + '_depa1403_' + str(self.getDocNumber(doc.name_real))

        if pdf_signed is not None:
            attachment = base64.decodebytes(pdf_signed)
            pdf_signed_name = filename_format + '.pdf'
            payload = MIMEBase('application', 'octate-stream')
            payload.set_payload(attachment)
            encoders.encode_base64(payload)

            payload.add_header('Content-Disposition', 'attachment', filename=pdf_signed_name)
            new_message.attach(payload)

        if len(doc.attachment_ids) > 0:
            for i, rec in enumerate(doc.attachment_ids):
                # print(self)
                attach_file_name = rec.name  # set filename
                if doc.name_real:
                    attach_file_name = filename_format + '_' + str(i + 1) + '.' + str(
                        self.getExtenstionName(rec.name))  # set filename
                attach_file = base64.decodebytes(rec.datas)
                # attach_file_name = os.path.dirname(os.path.abspath(__file__)) + '/dSign/sign_test.png'
                # attach_file = open(attach_file_name, 'rb')  # Open the file as binary mode
                payload = MIMEBase('application', 'octate-stream')
                payload.set_payload(attach_file)
                encoders.encode_base64(payload)  # encode the attachment

                # add payload header with filename
                payload.add_header('Content-Disposition', 'attachment', filename=attach_file_name)
                new_message.attach(payload)

        if len(doc.attached_file_in) > 0:
            for i, rec in enumerate(doc.attached_file_in):
                # print(self)
                attach_file_name = rec.name  # set filename
                if doc.name_real:
                    attach_file_name = filename_format + '_' + str(len(doc.attachment_ids) + i + 1) + '.' + str(
                        self.getExtenstionName(rec.datas_fname))  # set filename
                attach_file = base64.decodebytes(rec.datas)
                # attach_file_name = os.path.dirname(os.path.abspath(__file__)) + '/dSign/sign_test.png'
                # attach_file = open(attach_file_name, 'rb')  # Open the file as binary mode
                payload = MIMEBase('application', 'octate-stream')
                payload.set_payload(attach_file)
                encoders.encode_base64(payload)  # encode the attachment

                # add payload header with filename
                payload.add_header('Content-Disposition', 'attachment', filename=attach_file_name)
                new_message.attach(payload)

        imap_server.append('[Gmail]/Drafts', '', imaplib.Time2Internaldate(time.time()),
                           str(new_message).encode('utf-8'))
        imap_server.close()

    def set_default_email(self, doc_id, pdf_signed=None):
        # print(self)
        doc = self.env['document.internal.main'].search([('id', '=', [doc_id])])
        doc_speed = dict((
            ("1", "ปกติ"),
            ("2", "ด่วน"),
            ("3", "ด่วนมาก"),
            ("4", "ด่วนที่สุด")
        ))

        prefix_sub = ''
        if doc.speed != "1":
            prefix_sub = f'{doc_speed[doc.speed]} : '

        # content body
        subject = f"{prefix_sub}{doc.internal_title}"
        msgBody = "<br>".join(doc.material.split("\n"))
        for_doc = f"จึงเรียนมา{doc.for_document}"
        if doc.for_document == 'other':
            for_doc = doc.for_document_other
        suffix = ""
        if doc.suffix:
            suffix = doc.suffix
        msgText = f"""
                <html>
                  <head></head>
                  <body>
                    <p>
                        เรียน {doc.dear_select}<br><br>
                        {msgBody}<br><br>
                        {for_doc}<br><br>

                        <pre>สำนักงานส่งเสริมเศรษฐกิจดิจิทัล<br>{suffix}</pre><br><br>
                        ___________________________________________________________________________<br>
                        หากท่านได้รับอีเมลนี้แล้วกรุณาแจ้งการได้รับกลับมายังที่อยู่อีเมลนี้ด้วย จะขอบคุณยิ่ง<br><br>

                        อีเมล (และ/หรือเอกสารแนบ) นี้เป็นข้อมูลที่ความลับและอาจเป็นข้อมูลที่เป็นเอกสิทธิ์เฉพาะบุคคล การนำข้อมูลดังกล่าวไปใช้หรือเปิดเผยให้บุคคลอื่นใดล่วงรู้ เป็นการกระทำที่ไม่ได้รับอนุญาต หากท่านมิได้เป็นบุคคลที่อีเมลฉบับนี้ระบุถึงแล้ว กรุณาลบอีเมลนี้ออกจากคอมพิวเตอร์ที่ท่านได้รับ ทั้งนี้ ขอขอบคุณหากท่านได้แจ้งผู้ส่งถึงการจัดส่งอีเมลผิดพลาดด้วย<br><br>

                        This e-mail (and/or attachments) is confidential and may be privileged. Use or disclosure of it by anyone other than a designated addressee is unauthorized. If you are not an intended recipient, please delete this e-mail from the computer on which you received it. We thank you for notifying us immediately.
                    </p>
                  </body>
                </html>
                """

        # attachment and digital signature pdf file
        default_attachment_ids = []
        filename_format = 'สำเนาหนังสือลงนาม (Digital Signature)'
        type_lists = ['ประกาศ', 'ระเบียบ', 'ข้อบังคับ', 'คำสั่ง']
        if doc.name_real:
            reformat = '_depa'
            if 'สศด.' in doc.name_real:
                reformat = f'_depa{doc.name_real[4:].replace(" /", "_")}'
            elif bool([ele for ele in type_lists if (ele in doc.name_real)]):
                reformat = f'_depa_{doc.name_real.split("/")[0].replace(" ","_")}'
            filename_format = str(date.today().year + 543) + reformat

        if pdf_signed is not None:
            attach_vals = {
                'name': '%s.pdf' % filename_format,
                'datas': pdf_signed.datas,
                'datas_fname': '%s.pdf' % filename_format
            }
            pdf_signed_attachment = self.env['ir.attachment'].create(attach_vals)
            default_attachment_ids.append(pdf_signed_attachment.id)

        if len(doc.attachment_ids) > 0:
            for i, rec in enumerate(doc.attachment_ids):
                attach_file_name = rec.name  # set filename
                if doc.name_real:
                    attach_file_name = filename_format + '_' + str(i + 1) + '.' + str(
                        self.getExtenstionName(rec.name))  # set filename

                attach_vals = {
                    'name': attach_file_name,
                    'datas': rec.datas,
                    'datas_fname': attach_file_name
                }
                file_attachment = self.env['ir.attachment'].create(attach_vals)
                default_attachment_ids.append(file_attachment.id)

        if len(doc.attached_file_in) > 0:
            for i, rec in enumerate(doc.attached_file_in):
                attach_file_name = rec.name  # set filename
                if doc.name_real:
                    attach_file_name = filename_format + '_' + str(len(doc.attachment_ids) + i + 1) + '.' + str(
                        self.getExtenstionName(rec.datas_fname))  # set filename

                attach_vals = {
                    'name': attach_file_name,
                    'datas': rec.datas,
                    'datas_fname': attach_file_name
                }
                file_attachment = self.env['ir.attachment'].create(attach_vals)
                default_attachment_ids.append(file_attachment.id)

        self.env['saraban_send_email'].create({
            'name': subject,
            'body': msgText,
            'attachment_ids': [(6, 0, default_attachment_ids)],
            'email_owner_id': doc.create_uid.id,
            'destination_email': doc.sign_email,
            'real_name': doc.name_real,
            'sequence_name': doc.name
        })

    @api.onchange('action_request_otp_button')
    def send_mail_otp(self):
        # r = request
        # request.session['dsign_auth'] = False
        # print(self)
        if self.sign_otp_temp == str(TOTP.now()):
            time_remaining = TOTP.interval - datetime.now().timestamp() % TOTP.interval
            return {
                'warning': {
                    'title': 'แจ้งเตือนผู้ใช้งาน',
                    'message': f'ส่งรหัส OTP ไปยัง email ของคุณแล้ว \n กรุณาตรวจสอบหรือรอ 5 นาทีเพื่อส่งขอรหัสอีกครั้ง'}
            }

        if not self.sign_otp_temp:
            self.sign_otp_temp = '-'
        elif self.sign_otp_temp and self.sign_otp_temp != TOTP.now():
            # sent otp to approver email
            employee = self.env['hr.employee'].search(
                [
                    ('user_id', '=', self.env.uid)
                ],
                limit=1
            )
            email_to = employee.work_email
            self.sign_otp_temp = TOTP.now()
            # credential = os.path.dirname(os.path.abspath(__file__)) + '/dSign/credential_sarabun.txt'
            smtp_ssl_host = "smtp.gmail.com"
            smtp_ssl_port = 465
            hasSettings = self.env['depa_signature_setting'].search([], limit=1)
            if not hasSettings:
                raise ValidationError(
                    _('กรุณาตั้งค่า email สำหรับฉบับร่าง'))
            email_address = hasSettings.email_to_draft
            email_password = base64.b64decode(hasSettings.email_credential.encode('utf-8')).decode('utf-8')
            # email_address = "saraban@depa.or.th"
            # with open(credential, "r") as f:
            #     email_password = f.read().strip()
            print("OTP: ", self.sign_otp_temp)
            reply = f"""
                รหัส OTP ของท่านคือ {self.sign_otp_temp}
                กรุณาเก็บเป็นความลับ รหัสจะหมดอายุภายใน 5 นาที

                ขอบคุณครับ
                """
            msg = MIMEText(reply)
            msg['Subject'] = 'แจ้งรหัส OTP สำหรับการยืนยันตัวตนในการลงนาม Digital Signature'
            msg['From'] = email_address
            msg['To'] = email_to
            server = smtplib.SMTP_SSL(smtp_ssl_host, smtp_ssl_port)
            server.login(email_address, email_password)
            # server.starttls()
            server.sendmail(email_address, email_to, msg.as_string())
            server.quit()

            hasSettingsLines = self.env['depa_signature_setting_lines'].search(
                [
                    ('employee_id.user_id', '=', self.env.uid),
                    ('disable', '=', False),
                    ('active', '=', True)
                ], limit=1)
            if hasSettingsLines:
                if hasSettingsLines.otp_default:
                    self.sign_otp = self.sign_otp_temp

            print('sent otp to approver email.')
            return {
                'warning': {
                    'title': 'แจ้งเตือนผู้ใช้งาน',
                    'message': f'ส่งรหัส OTP ไปยัง email ของคุณแล้ว \n กรุณาตรวจสอบที่ {email_to}'}
            }

    def _get_current_time(self):
        now = datetime.now(timezone('Asia/Bangkok')).strftime("%H%M")
        return now

    def _get_weekday(self):
        today = datetime.now(timezone('Asia/Bangkok')).date().weekday()
        return today

    def _default_fiscal_year(self):
        fiscal_year_obj = self.env['fw_pfb_fin_system_fiscal_year'].search([
            ('date_start', '<=', date.today()),
            ('date_end', '>=', date.today()),
        ], limit=1)

        if fiscal_year_obj:
            return fiscal_year_obj.id

    def _check_in_holiday(self):
        this_year_holidays = self.env['leave_request_public_holidays'].search([
            ('fiscal_year_id', '=', self._default_fiscal_year())
        ], limit=1)

        if this_year_holidays:
            public_holidays = self.env['leave_request_public_holidays_lines'].search([
                ('leave_request_public_holidays_lines_id', '=', this_year_holidays.id)
            ])
            if public_holidays:
                data = []
                for holiday in public_holidays:
                    holiday_detail = self.env['leave_request_public_holidays_lines'].search([
                        ('id', '=', holiday['id'])
                    ])
                    if holiday_detail:
                        if date.today() == holiday_detail['date']:
                            return True
        return False

    @api.multi
    def action_approve_signature(self):
        super(MakeApprovalSignatureWizardInherit, self).action_approve()
        setting_id = self.env['document.internal.main'].browse(int(self.setting_id))
        is_sign_group = self.env.user.has_group('depa_signature.group_user_digital_signature')

        if setting_id.document_type in ['ประกาศพัสดุ']:
            string = "เนื่องจากการอนุมัติงานพัสดุมีผลกับระบบ eGP\nทั้งนี้ได้เลยระยะเวลาทำการ\nจึงขอให้อนุมัติงานอีกครั้งในวันทำการถัดไป เวลา 08.00-17.00 น."
            
            if (not 800 <= int(self._get_current_time()) <= 1700) or (self._get_weekday() in [5,6]) or (self._check_in_holiday()):
                raise ValidationError(string)

        if setting_id.document_type in ['หนังสือภายนอก+หนังสือรับรอง', 'ประกาศ', 'ระเบียบ', 'ข้อบังคับ', 'หนังสือรับรอง'] or 'คำสั่ง' in setting_id.document_type:
            if (setting_id.circular_letter_check == True) and setting_id.document_type == 'หนังสือภายนอก+หนังสือรับรอง':
                if setting_id.name_real:
                    name_real = setting_id.name_real
                    code = name_real.split('/')[0]
                    number = name_real.split('/')[1]
                    doc_name_real = code + "/ว." + number

                    setting_id.update({
                        'name_real': doc_name_real,
                        'date_document_real': fields.Date.today()
                    })

            if setting_id.state == 'done' and is_sign_group:
                cad_pass = ''
                # Check 2FA IF d-Sign
                if self.sign_otp:
                    if not TOTP.verify(self.sign_otp):
                        raise ValidationError(
                            _('รหัส OTP ไม่ถูกต้องหรือไม่ได้ใช้งานภายใน 5 นาที \n กรุณากดขอรหัส OTP ใหม่อีกครั้ง'))
                    else:
                        # create session for 2FA authentication
                        self.sign_otp = False
                        self.sign_otp_temp = False
                        request.session['dsign_auth'] = True

                # Call Sign Document
                if self.sign_method == 'sign':
                    sign_data = self.sign
                else:
                    sign_data = self.img

                hasSettingsLines = self.env['depa_signature_setting_lines'].search(
                    [
                        ('employee_id.user_id', '=', self.env.uid),
                        ('disable', '=', False),
                        ('active', '=', True)
                    ], limit=1)
                if not hasSettingsLines:
                    raise ValidationError(
                        _('กรุณาตั้งค่า Digital Signature ไฟล์ให้เสร็จสมบูรณ์'))
                cad_pass = hasSettingsLines.cad_password
                # print("CAD: ", cad_pass)

                send_success = False
                # ! กรณีเลือกหนังสือเวียน
                if setting_id.circular_letter_check and setting_id.document_type == 'หนังสือภายนอก+หนังสือรับรอง':
                    lists = setting_id.invitation_lines_ids
                    names = []
                    doc_ids = []

                    converted_document = self.convert_plain_document(base64.b64decode(setting_id.file_for_signature))
                    for invitation in lists:
                        sign_status, data_pdf = self.sign_document(int(setting_id.id), sign_data, cad_pass,
                                                                   invitation_name=invitation.invitation_name,
                                                                   converted_document=converted_document)
                        if sign_status:
                            invitation_line = self.env['invitation_lines'].browse(int(invitation.id))
                            attach_vals = {
                                'name': '%s.pdf' % ('DocumentSinged_' + datetime.now(timezone('Asia/Bangkok')).strftime(
                                    "%Y%m%d %H%M%S")),
                                'datas': data_pdf,
                                'datas_fname': '%s.pdf' % (
                                            'หนังสือเวียน_' + setting_id.name_real + '_' + invitation.invitation_name)
                            }
                            doc_id = self.env['ir.attachment'].create(attach_vals)
                            if doc_id:
                                doc_ids.append(doc_id.id)
                                invitation_line.update({
                                    'document_pdf_id': int(doc_id.id)
                                })
                    if doc_ids:
                        setting_id.update({
                            'doc_pdf_signed': [(6, 0, doc_ids)],
                            'head_officer_digital_signed': True
                        })
                        send_success = True
                    self.set_default_email(int(setting_id.id))
                
                # ! กรณีเป็นคำสั่งประกาศ จะต้องมีการเซ็น 2 ครั้ง
                # elif setting_id.document_type == "คำสั่ง ข":
                # # elif setting_id.document_type == "ประกาศพัสดุ":
                #     signatures = [
                #         {'x': 520, 'y': 50, 'position': 'P3'},  # ผู้ตรวจสอบ
                #         {'x': 745, 'y': 50, 'position': 'P4'}   # ผู้อนุมัติ
                #     ]
                #     pdf_file = setting_id.file_for_signature
                #     for signature in signatures:
                #         employee_signature, employee_cad_pass = self.get_account_approver_sign(position=signature['position'])
                #         sign_status, data_pdf = self.sign_document(int(setting_id.id), employee_signature, employee_cad_pass,
                #                                 specific_x=signature['x'], specific_y=signature['y'],
                #                                 pdf_file=pdf_file
                #         )
                #         pdf_file = data_pdf
                #     if sign_status:
                #         attach_vals = {
                #             'name': '%s.pdf' % ('DocumentSinged_' + datetime.now(timezone('Asia/Bangkok')).strftime(
                #                 "%Y%m%d %H%M%S")),
                #             'datas': data_pdf,
                #             'datas_fname': '%s.pdf' % (
                #                         'DocumentSinged_' + datetime.now(timezone('Asia/Bangkok')).strftime(
                #                     "%Y%m%d %H%M%S"))
                #         }
                #         doc_id = self.env['ir.attachment'].create(attach_vals)
                #         if doc_id:
                #             setting_id.update({
                #                 'doc_pdf_signed': [(6, 0, [doc_id.id])],
                #                 'head_officer_digital_signed': True
                #             })
                #         self.set_default_email(int(setting_id.id), doc_id)
                
                # ! กรณีอื่นๆ จะสร้างหนังสือใหม่แค่รอบเดียว
                else:  
                    sign_status, data_pdf = self.sign_document(int(setting_id.id), sign_data, cad_pass)
                    # print("Sign: ",sign_status)

                    if sign_status:
                        # Save d-sign doc into attachment
                        attach_vals = {
                            'name': '%s.pdf' % ('DocumentSinged_' + datetime.now(timezone('Asia/Bangkok')).strftime(
                                "%Y%m%d %H%M%S")),
                            'datas': data_pdf,
                            'datas_fname': '%s.pdf' % (
                                        'DocumentSinged_' + datetime.now(timezone('Asia/Bangkok')).strftime(
                                    "%Y%m%d %H%M%S"))
                        }
                        doc_id = self.env['ir.attachment'].create(attach_vals)
                        if doc_id:
                            setting_id.update({
                                'doc_pdf_signed': [(6, 0, [doc_id.id])],
                                'head_officer_digital_signed': True
                            })
                            send_success = True

                        # set draft email
                        # self.set_draft_email(int(setting_id.id), doc_id.datas)
                        self.set_default_email(int(setting_id.id), doc_id)

                        # return {
                        #     'type': 'ir.actions.act_url',
                        #     'url': 'web/content/%s?download=true' % (doc_id.id),
                        #     'target': 'current',
                        # }

                    # End Sign Document

                # เมื่อเซ็นหนังสือเสร็จสมบูรณ์ได้แล้ว จะให้ทำการส่ง LINE เพื่อแจ้งเตือนไปยังผู้สร้าง
                if send_success:
                    token_line = self.env['hr.employee'].search([('user_id', "=", setting_id.create_uid.id)], limit=1).token_line

                    if token_line:
                        base_url = request.httprequest.url_root
                        url = f"{base_url}web#id={setting_id.id}&model=document.internal.main&view_type=form"

                        string = f"\nแจ้งเตือนหนังสือเสร็จสมบูรณ์\n" \
                        f"เลขหนังสือ: {setting_id.name}\n" \
                        f"เรื่อง: {setting_id.subject}\n" \
                        f"ถึง: {setting_id.create_uid.name}\n" \
                        f"URL: {url}\n" \

                        payload = {
                            "message": string
                        }

                        try:
                            r = requests.post(
                                "https://notify-api.line.me/api/notify",
                                headers={"Authorization": "Bearer {}".format(token_line)},
                                params = payload
                            )
                        except Exception as e:
                            raise ValueError(e)
